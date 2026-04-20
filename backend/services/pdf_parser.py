import logging
import os

import PyPDF2
import pdfplumber

logger = logging.getLogger(__name__)


def _extract_pdf_pdfplumber(file_path: str) -> str:
    """Advanced PDF extraction — better with tables and multi-column layouts."""
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as exc:
        logger.warning("pdfplumber failed on %s: %s", file_path, exc)
    return text.strip()


def _extract_pdf_pypdf2(file_path: str) -> str:
    """Fallback PDF extraction."""
    text = ""
    try:
        with open(file_path, "rb") as fh:
            reader = PyPDF2.PdfReader(fh)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as exc:
        logger.warning("PyPDF2 failed on %s: %s", file_path, exc)
    return text.strip()


def _extract_docx(file_path: str) -> str:
    """Extract text from a .docx file using python-docx."""
    try:
        import docx  # python-docx
        doc = docx.Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs).strip()
    except Exception as exc:
        logger.error("DOCX extraction failed on %s: %s", file_path, exc)
        return ""


def parse_resume(file_path: str) -> str:
    """
    Extract plain text from a resume file.
    Supports PDF (tries pdfplumber, falls back to PyPDF2) and DOCX.
    Returns empty string if extraction fails.
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext in (".docx", ".doc"):
        text = _extract_docx(file_path)
        if not text:
            logger.error("Could not extract text from DOCX: %s", file_path)
        return text

    # Default: PDF
    text = _extract_pdf_pdfplumber(file_path)
    if not text or len(text) < 50:
        logger.info("pdfplumber returned short text, falling back to PyPDF2 for %s", file_path)
        text = _extract_pdf_pypdf2(file_path)

    if not text:
        logger.error("All PDF extraction methods failed for %s", file_path)

    return text
